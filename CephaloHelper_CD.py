from PIL import ImageTk, Image
import tkinter as tk
import helper as helper

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION

'''
Points:

[L0_sella,
L1_nasion,
L2_porion,
L3_orbitale,
L4_ENP,
L5_ENA,
L6_point_A,
L7_point_B,
L8_pogonion,
L9_menton,
L10_gnathion,
L11_symphyse
L12_gonion,
L13_articulate,
L14_I1,
L15_I2,
L16_i1,
L17_i2,
L18_M1,
L19_M2,
L20_m1,
L21_m2,
L22_POA,
L23_POP]
'''

class Cephalo(tk.Tk):
    def __init__(self):
        tk.Tk.__init__(self)
        self.name = "Bougheda Fady"
        self.x = self.y = 0
        self.pts = []
        self.pts_text = ["S", "Na", "Po", "Or", "ENP", "ENA", "A", "B", "Pog", "Me", "Gn", "D", "Go",
                         "Ar", "I1", "I2", "i_1", "i_2", "M1", "M2", "m_1", "m_2", "POA", "POP", "clb1", "clb2"]

        self.angles_text = ["SNA", "SNB", "ANB", "Ao-Bo", "SND", "Â.Conv", "Â.Facial", "FMA", "Â.Gon", "Axe Y",
                            "Etage Superieur", "Etage Inferieur", "I/F", "i/m", "I/i", "Alpha", "Beta"]

        self.canvas = tk.Canvas(self, width=1000, height=750, cursor="cross")
        self.canvas.pack(side="top", fill="both", expand=True)
        self.img = ImageTk.PhotoImage(Image.open(self.name+".jpg"))
        # self.img = ImageTk.PhotoImage(Image.open("from_phone.jpg"))
        self.canvas.create_image(0, 0, image=self.img, anchor='nw')
        self.canvas.bind("<ButtonPress-1>", self.on_button_press)
        self.canvas.bind("<ButtonPress-3>", self.on_back)

    def on_back(self, event):
        print(self.pts_text[len(self.pts) - 1])
        self.canvas.delete(self.pts_text[len(self.pts) - 1])
        self.canvas.delete(str(self.pts_text[len(self.pts) - 1]) + "_pt")
        self.pts.pop(-1)

    def calculate(self):
        SNA = helper.get_angle(self.pts[0], self.pts[1], self.pts[6])
        SNA_dsc = helper.SNA_dsc
        if SNA>86:
            SNA_dsc = helper.SNA_dsc_plus
        elif SNA < 80:
            SNA_dsc = helper.SNA_dsc_minus

        SNB = helper.get_angle(self.pts[0], self.pts[1], self.pts[7])
        SNB_dsc = helper.SNB_dsc
        if SNB > 83:
            SNB_dsc = helper.SNB_dsc_plus
        elif SNB < 77:
            SNB_dsc = helper.SNB_dsc_minus

        ANB = round(SNA - SNB)
        ANB_dsc = helper.ANB_dsc
        if ANB > 4:
            ANB_dsc = helper.ANB_dsc_plus
        elif ANB < 2:
            ANB_dsc = helper.ANB_dsc_minus

        AOBO = helper.aobo_verify(self.pts[6], self.pts[7], self.pts[22], self.pts[23], self.pts[-1], self.pts[-2])
        AOBO_dsc = helper.AoBo_dsc
        if AOBO > 3:
            AOBO_dsc = helper.AoBo_dsc_plus
        elif AOBO < -1:
            AOBO_dsc = helper.AoBo_dsc_minus

        SND = helper.get_angle(self.pts[0], self.pts[1], self.pts[11])
        SND_dsc = helper.SND_dsc
        if SND > 76:
            SND_dsc = helper.SND_dsc_plus
        elif SND < 76:
            SND_dsc = helper.SND_dsc_minus

        AC = round(180 - helper.get_angle(self.pts[8], self.pts[6], self.pts[1]))
        AC_dsc = helper.AC_dsc
        if AC > 11:
            AC_dsc = helper.AC_dsc_plus
        elif AC < 1:
            AC_dsc = helper.AC_dsc_minus

        AF = helper.get_angle(self.pts[2], self.pts[3], self.pts[8])
        AF_dsc = helper.AF_dsc
        if AF > 93:
            AF_dsc = helper.AF_dsc_plus
        elif AF < 87:
            AF_dsc = helper.AF_dsc_minus

        FMA = helper.get_angle_lines(self.pts[2], self.pts[3], self.pts[12], self.pts[9])
        FMA_dsc = helper.FMA_dsc
        if FMA > 31:
            FMA_dsc = helper.FMA_dsc_plus
        elif FMA < 23:
            FMA_dsc = helper.FMA_dsc_minus

        AG = helper.get_angle(self.pts[9], self.pts[12], self.pts[13])
        AG_dsc = helper.AG_dsc
        if AG > 134:
            AG_dsc = helper.AG_dsc_plus
        elif AG < 122:
            AG_dsc = helper.AG_dsc_minus

        AXE_Y = helper.get_angle_lines(self.pts[0], self.pts[10], self.pts[2], self.pts[3])
        AXE_Y_dsc = helper.AXE_Y_dsc
        if AXE_Y > 62:
            AXE_Y_dsc = helper.AXE_Y_dsc_plus
        elif AXE_Y < 56:
            AXE_Y_dsc = helper.AXE_Y_dsc_minus

        E_SUP, E_INF = helper.rapport_etage(self.pts[5], self.pts[1], self.pts[9])
        E_SUP_dsc = helper.E_SUP_dsc
        if E_SUP > 47.5:
            E_SUP_dsc = helper.E_SUP_dsc_plus
        elif E_SUP < 43.5:
            E_SUP_dsc = helper.E_SUP_dsc_minus

        E_INF_dsc = helper.E_INF_dsc
        if E_INF > 56.5:
            E_INF_dsc = helper.E_INF_dsc_plus
        elif E_INF < 52.5:
            E_INF_dsc = helper.E_INF_dsc_minus

        I_F = helper.get_angle_lines(self.pts[14], self.pts[15], self.pts[2], self.pts[3])
        I_F_dsc = helper.I_F_dsc
        if I_F > 110:
            I_F_dsc = helper.I_F_dsc_plus
        elif I_F < 104:
            I_F_dsc = helper.I_F_dsc_minus

        I_M = helper.get_angle_lines(self.pts[16], self.pts[17], self.pts[12], self.pts[9])
        I_M_dsc = helper.I_M_dsc
        if I_M > 93:
            I_M_dsc = helper.I_M_dsc_plus
        elif I_M < 87:
            I_M_dsc = helper.I_M_dsc_minus

        I_I = helper.get_angle_lines(self.pts[14], self.pts[15], self.pts[16], self.pts[17])
        I_I_dsc = helper.I_I_dsc
        if I_I > 140:
            I_I_dsc = helper.I_I_dsc_plus
        elif I_I < 130:
            I_I_dsc = helper.I_I_dsc_minus

        ALPHA = helper.get_angle_lines(self.pts[18], self.pts[19], self.pts[23], self.pts[22])
        ALPHA_dsc = helper.ALPHA_dsc
        if ALPHA > 93:
            ALPHA_dsc = helper.ALPHA_dsc_plus
        elif ALPHA < 87:
            ALPHA_dsc = helper.ALPHA_dsc_minus

        BETA = helper.get_angle_lines(self.pts[20], self.pts[21], self.pts[23], self.pts[22])
        BETA_dsc = helper.BETA_dsc
        if BETA > 103:
            BETA_dsc = helper.BETA_dsc_plus
        elif BETA < 97:
            BETA_dsc = helper.BETA_dsc_minus


        self.angles = [SNA, SNB, ANB, AOBO, SND, AC, AF, FMA, AG,
                       AXE_Y, E_SUP, E_INF, I_F, I_M, I_I, ALPHA, BETA]

        self.description = [SNA_dsc, SNB_dsc, ANB_dsc, AOBO_dsc, SND_dsc, AC_dsc,
                            AF_dsc, FMA_dsc, AG_dsc, AXE_Y_dsc, E_SUP_dsc, E_INF_dsc,
                            I_F_dsc, I_M_dsc, I_I_dsc, ALPHA_dsc, BETA_dsc]

# //////////////////////////////////
        document = Document()
        # section = document.sections[-1]
        # new_width, new_height = section.page_height, section.page_width
        # new_section = document.add_section(WD_SECTION.NEW_PAGE)
        # new_section.orientation = WD_ORIENT.LANDSCAPE
        # new_section.page_width = new_width
        # new_section.page_height = new_height

        document.add_heading(self.name, 0)
# ////////////////////////////////

        for i, angle in enumerate(self.angles):
            self.canvas.create_text(120, 20 * i + 500, fill="black", font="Times 12 bold",
                                    text=self.angles_text[i] + ": " + str(angle))
            p = self.angles_text[i] + ": " + str(angle) + ", " + self.description[i]
            document.add_paragraph(p, style='List Bullet')
        document.add_picture(self.name + '.jpg', width=Inches(4))
        document.save(self.name + ".docx")

    # def calculate(self):
    #
    #     AOBO = helper.aobo_verify(self.pts[0], self.pts[1], self.pts[2], self.pts[3], self.pts[-1], self.pts[-2])
    #     self.canvas.create_text(120, 20*1+500, fill="black", font="Times 12 bold",
    #                             text=str(AOBO))


    def on_button_press(self, event):
        # 26
        if len(self.pts) == 26:
            self.calculate()
        else:
            self.x = event.x
            self.y = event.y
            x0, y0 = (self.x - 2, self.y - 2)
            x1, y1 = (event.x + 2, event.y + 2)
            self.canvas.create_oval(x0, y0, x1, y1, fill="red", tag=str(self.pts_text[len(self.pts)]) + "_pt")
            self.canvas.create_text(self.x, self.y - 12, fill="red", font="Times 12 bold",
                                    text=self.pts_text[len(self.pts)], tag=self.pts_text[len(self.pts)])
            self.pts.append([float(self.x), float(self.y)])

            # if len(self.pts) == 3:
            #     self.angle = helper.get_angle(self.pts[0], self.pts[1], self.pts[2])
            #     self.canvas.create_text(self.pts[1][0], self.pts[1][1]-12, fill="red", font="Times 12",
            #                              text=self.angle)


if __name__ == "__main__":
    app = Cephalo()
    app.mainloop()






